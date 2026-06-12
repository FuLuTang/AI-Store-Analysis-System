"""底层实现：在 workspace 内搜索文本、正则模式或聊天历史记录"""
import json
import re
import sqlite3
import tarfile
import zipfile
from pathlib import Path
from ...workspace import Workspace
from ...file_domains import DEFAULT_FILE_DOMAINS, join_domain_path, split_domain_path
CHAT_HISTORY_FILENAME = "chat.jsonl"
MAX_PREVIEW_CHARS = 1200
MAX_TEXT_LINES = 2000


def _history_time(message: dict) -> str:
    return str(message.get("datetime") or message.get("time") or "").strip()


def _public_history_records_with_reasoning(messages: list[dict]) -> list[dict]:
    visible: list[dict] = []
    for message in messages:
        role = str(message.get("role", "")).strip().lower()
        if role == "system":
            name = str(message.get("name", "")).strip().lower()
            if name == "notice":
                notice = {
                    "role": "notice",
                    "content": str(message.get("content") or ""),
                }
                history_time = _history_time(message)
                if history_time:
                    notice["datetime"] = history_time
                visible.append(notice)
                continue
            if str(message.get("title") or "").strip():
                card = {
                    "role": "card",
                    "name": str(message.get("name") or ""),
                    "title": str(message.get("title") or ""),
                    "detail": str(message.get("detail") or ""),
                }
                options = message.get("options")
                if isinstance(options, list):
                    card["options"] = [str(item) for item in options]
                if "choice" in message:
                    card["choice"] = str(message.get("choice") or "")
                history_time = _history_time(message)
                if history_time:
                    card["datetime"] = history_time
                visible.append(card)
            continue
        if role in {"user", "assistant"}:
            content_val = str(message.get("content") or "")
            reasoning_val = str(message.get("reasoning_content") or "")
            tool_calls = message.get("tool_calls")
            if role == "assistant" and not content_val.strip() and not reasoning_val.strip() and not tool_calls:
                continue
            public_msg = {
                "role": role,
                "content": content_val,
            }
            if reasoning_val.strip():
                public_msg["reasoning_content"] = reasoning_val
            history_time = _history_time(message)
            if history_time:
                public_msg["datetime"] = history_time
            attachments = message.get("attachments")
            if isinstance(attachments, list):
                public_msg["attachments"] = attachments
            visible.append(public_msg)
    return visible


def _read_chat_history_lines(ws: Workspace) -> tuple[str, list[dict]]:
    candidates = []
    if ws.has_multi_read_roots:
        candidates.append("chatbot/" + CHAT_HISTORY_FILENAME)
        candidates.append(CHAT_HISTORY_FILENAME)
    else:
        candidates.append(CHAT_HISTORY_FILENAME)

    history_path = None
    for candidate in candidates:
        try:
            resolved = ws.resolve_read(candidate)
        except ValueError:
            continue
        if resolved.exists() and resolved.is_file():
            history_path = resolved
            break

    if history_path is None:
        raise FileNotFoundError("未找到 chat_history 记录文件")

    messages: list[dict] = []
    with history_path.open("r", encoding="utf-8") as fh:
        for line in fh:
            raw = line.strip()
            if not raw:
                continue
            try:
                record = json.loads(raw)
            except json.JSONDecodeError:
                continue
            if isinstance(record, dict):
                messages.append(record)

    visible = _public_history_records_with_reasoning(messages)
    return str(history_path), visible


def _search_serialized_lines(lines: list[dict], pattern: str, regex: bool, max_matches: int, path_label: str) -> dict:
    rx = None
    pattern_lower = None
    if regex:
        rx = re.compile(pattern, re.IGNORECASE)
    else:
        pattern_lower = pattern.lower()

    results = []
    total_matches = 0
    hit_limit = False

    for line_idx, record in enumerate(lines, start=1):
        serialized = json.dumps(record, ensure_ascii=False)
        matched = False
        if regex and rx:
            matched = bool(rx.search(serialized))
        elif pattern_lower is not None:
            matched = pattern_lower in serialized.lower()
        if not matched:
            continue

        preview = serialized
        if len(preview) > 500:
            preview = preview[:500] + "..."
        results.append({
            "line": line_idx,
            "text": preview,
        })
        total_matches += 1
        if total_matches >= max_matches:
            hit_limit = True
            break

    output = {
        "pattern": pattern,
        "regex": regex,
        "domain": "chat_history",
        "path": path_label,
        "total_lines": len(lines),
        "total_matches_found": total_matches,
        "hit_limit": hit_limit,
        "results": [
            {
                "path": path_label,
                "total_matches": len(results),
                "matches": results,
            }
        ] if results else [],
    }
    if hit_limit:
        output["note"] = f"匹配数量已达到最大上限 {max_matches}，部分匹配未展示。"
    return output


def _is_hidden_relative(path: Path, root: Path) -> bool:
    try:
        rel_parts = path.relative_to(root).parts
    except ValueError:
        return True
    return any(part.startswith(".") for part in rel_parts)


def _domain_path(domain: str, root: Path, path: Path) -> str:
    rel = str(path.relative_to(root)).replace("\\", "/")
    return join_domain_path(domain, rel)


def _match_text(text: str, rx, pattern_lower: str | None) -> bool:
    if rx:
        return bool(rx.search(text))
    return bool(pattern_lower and pattern_lower in text.lower())


def _preview(text: str, limit: int = 500) -> str:
    text = " ".join(str(text or "").split())
    if len(text) <= limit:
        return text
    return text[:limit] + "..."


def _iter_text_matches(path: Path, rx, pattern_lower: str | None, remaining: int):
    try:
        with path.open("r", encoding="utf-8", errors="replace") as fh:
            for line_idx, line in enumerate(fh, start=1):
                if line_idx > MAX_TEXT_LINES or remaining <= 0:
                    break
                if _match_text(line, rx, pattern_lower):
                    yield {
                        "match_type": "content",
                        "line": line_idx,
                        "text": _preview(line),
                    }
                    remaining -= 1
    except Exception:
        return


def _extract_xlsx_preview(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        parts = []
        for sheet_name in wb.sheetnames[:5]:
            ws = wb[sheet_name]
            parts.append(f"sheet: {sheet_name}")
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if row_idx > 5:
                    break
                values = [str(value) for value in row if value is not None]
                if values:
                    parts.append(" | ".join(values[:20]))
        return "\n".join(parts)[:MAX_PREVIEW_CHARS]
    finally:
        wb.close()


def _extract_pdf_preview(path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page_idx, page in enumerate(pdf.pages[:3], start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"page {page_idx}: {text.strip()}")
    return "\n".join(parts)[:MAX_PREVIEW_CHARS]


def _extract_docx_preview(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs[:30]:
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else ""
            parts.append(f"{style}: {text}" if style else text)
    for table_idx, table in enumerate(doc.tables[:3], start=1):
        if table.rows:
            cells = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
            if cells:
                parts.append(f"table {table_idx}: {' | '.join(cells)}")
    return "\n".join(parts)[:MAX_PREVIEW_CHARS]


def _extract_sqlite_preview(path: Path) -> str:
    conn = sqlite3.connect(f"file:{path.resolve()}?mode=ro", uri=True)
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
        parts = []
        for (table_name,) in cursor.fetchall()[:20]:
            cursor.execute(f"PRAGMA table_info(\"{table_name}\")")
            cols = [row[1] for row in cursor.fetchall()]
            parts.append(f"table: {table_name} columns: {', '.join(cols)}")
        return "\n".join(parts)[:MAX_PREVIEW_CHARS]
    finally:
        conn.close()


def _extract_archive_preview(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".zip":
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()[:100]
        return "\n".join(names)[:MAX_PREVIEW_CHARS]
    if ext in {".tar", ".gz"}:
        with tarfile.open(path) as archive:
            names = archive.getnames()[:100]
        return "\n".join(names)[:MAX_PREVIEW_CHARS]
    if ext == ".rar":
        import rarfile
        with rarfile.RarFile(path) as archive:
            names = archive.namelist()[:100]
        return "\n".join(names)[:MAX_PREVIEW_CHARS]
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()[:100]
    return "\n".join(names)[:MAX_PREVIEW_CHARS]


def _structured_preview(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext in {".txt", ".log", ".md", ".py", ".csv", ".json", ".jsonl"}:
        return "text", ""
    if ext in {".xlsx", ".xls"}:
        return "excel_preview", _extract_xlsx_preview(path)
    if ext == ".pdf":
        return "pdf_preview", _extract_pdf_preview(path)
    if ext in {".docx", ".doc"}:
        return "word_preview", _extract_docx_preview(path)
    if ext in {".sqlite", ".db"}:
        return "sqlite_preview", _extract_sqlite_preview(path)
    if ext in {".zip", ".tar", ".gz", ".rar"}:
        return "archive_preview", _extract_archive_preview(path)
    return "unsupported", ""


def _collect_domain_items(domain: str, root: Path, target: Path) -> tuple[list[Path], list[Path]]:
    directories = []
    files = []
    candidates = [target]
    if target.is_dir():
        candidates.extend(sorted(target.rglob("*")))
    for p in candidates:
        if _is_hidden_relative(p, root):
            continue
        if p.is_dir():
            directories.append(p)
        elif p.is_file():
            files.append(p)
    return directories, files


def _resolve_search_domains(ws: Workspace, domain: str, path: str | None):
    read_roots = ws.read_roots
    if domain not in {"", "all", "chat_history"} and path is None:
        path = domain

    if path:
        target = ws.resolve_read(path)
        if ws.has_multi_read_roots:
            domain_name, _ = split_domain_path(path, allowed_domains=read_roots.keys())
        else:
            domain_name = DEFAULT_FILE_DOMAINS[0]
        return [(domain_name, read_roots[domain_name].resolve(), target)]

    return [
        (domain_name, root.resolve(), root.resolve())
        for domain_name, root in read_roots.items()
    ]


def search_files_impl(
    ws: Workspace,
    pattern: str,
    domain: str = "all",
    path: str = None,
    regex: bool = False,
    max_matches: int = 50,
) -> str:
    """在工作区内检索匹配的文本，返回匹配行和行号的 JSON。"""
    max_matches = max(1, min(int(max_matches or 50), 200))
    domain = str(domain or "all").strip()

    if domain == "chat_history":
        try:
            history_path, visible_records = _read_chat_history_lines(ws)
        except FileNotFoundError as e:
            return json.dumps({"error": str(e), "domain": domain}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": f"读取聊天历史失败: {e}", "domain": domain}, ensure_ascii=False)

        try:
            output = _search_serialized_lines(visible_records, pattern, regex, max_matches, history_path)
        except re.error as e:
            return json.dumps({"error": f"无效的正则表达式: {e}"}, ensure_ascii=False)
        return json.dumps(output, ensure_ascii=False)

    # 编译正则或转换子串模式
    rx = None
    pattern_lower = None
    if regex:
        try:
            rx = re.compile(pattern, re.IGNORECASE)
        except re.error as e:
            return json.dumps({"error": f"无效的正则表达式: {e}"}, ensure_ascii=False)
    else:
        pattern_lower = pattern.lower()

    try:
        search_domains = _resolve_search_domains(ws, domain, path)
    except ValueError as e:
        return json.dumps({"error": str(e), "path": path or domain}, ensure_ascii=False)

    domains_output = []
    total_matches = 0
    total_files = 0
    total_dirs = 0

    for domain_name, root, target in search_domains:
        if not target.exists():
            domains_output.append({
                "domain": domain_name,
                "root_path": _domain_path(domain_name, root, target),
                "root_absolute_path": str(target.resolve()),
                "error": f"路径不存在: {_domain_path(domain_name, root, target)}",
                "results": [],
                "totals": {
                    "files_searched": 0,
                    "directories_searched": 0,
                    "matches_found": 0,
                    "hit_limit": False,
                },
            })
            continue

        directories, files = _collect_domain_items(domain_name, root, target)
        domain_results = []
        domain_matches = 0
        hit_limit = False

        for directory in directories:
            if domain_matches >= max_matches:
                hit_limit = True
                break
            rel_path = _domain_path(domain_name, root, directory)
            if _match_text(directory.name, rx, pattern_lower):
                domain_results.append({
                    "path": rel_path,
                    "absolute_path": str(directory.resolve()),
                    "type": "directory",
                    "total_matches": 1,
                    "matches": [{
                        "match_type": "directory_name",
                        "text": directory.name,
                    }],
                })
                domain_matches += 1

        for file_path in files:
            if domain_matches >= max_matches:
                hit_limit = True
                break

            rel_path = _domain_path(domain_name, root, file_path)
            file_matches = []
            if _match_text(file_path.name, rx, pattern_lower):
                file_matches.append({
                    "match_type": "file_name",
                    "text": file_path.name,
                })

            remaining = max_matches - domain_matches - len(file_matches)
            if remaining > 0:
                try:
                    preview_kind, preview_text = _structured_preview(file_path)
                    if preview_kind == "text":
                        file_matches.extend(
                            _iter_text_matches(file_path, rx, pattern_lower, remaining)
                        )
                    elif preview_text and _match_text(preview_text, rx, pattern_lower):
                        file_matches.append({
                            "match_type": "content",
                            "source": preview_kind,
                            "text": _preview(preview_text),
                        })
                except Exception:
                    # Unreadable structured files still participate through file-name matches.
                    pass

            if file_matches:
                if domain_matches + len(file_matches) > max_matches:
                    file_matches = file_matches[:max(0, max_matches - domain_matches)]
                    hit_limit = True
                domain_results.append({
                    "path": rel_path,
                    "absolute_path": str(file_path.resolve()),
                    "type": "file",
                    "total_matches": len(file_matches),
                    "matches": file_matches,
            })
                domain_matches += len(file_matches)

        total_files += len(files)
        total_dirs += len(directories)
        total_matches += domain_matches
        domains_output.append({
            "domain": domain_name,
            "root_path": _domain_path(domain_name, root, target),
            "root_absolute_path": str(target.resolve()),
            "results": domain_results,
            "totals": {
                "files_searched": len(files),
                "directories_searched": len(directories),
                "matches_found": domain_matches,
                "hit_limit": hit_limit,
            },
        })

    output = {
        "pattern": pattern,
        "regex": regex,
        "domain": domain,
        "domains": [],
        "total_files_searched": 0,
        "total_directories_searched": 0,
        "total_matches_found": 0,
        "hit_limit": False,
    }

    output["domains"] = domains_output
    output["total_files_searched"] = total_files
    output["total_directories_searched"] = total_dirs
    output["total_matches_found"] = total_matches
    output["hit_limit"] = any(
        item.get("totals", {}).get("hit_limit") for item in domains_output
    )

    if output["hit_limit"]:
        output["note"] = f"单个域匹配数量达到最大上限 {max_matches} 时会截断该域后续匹配。"
        
    return json.dumps(output, ensure_ascii=False)
