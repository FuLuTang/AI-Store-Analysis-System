"""底层实现：解析并探测工作区内各种文档的结构信息 (XLSX, CSV, PDF, DOCX, JSON, MD, TXT, ZIP, SQLite)"""
import json
import io
import re
import zipfile
import sqlite3
from pathlib import Path
from urllib.request import pathname2url
from typing import Callable, Optional
from ...workspace import Workspace


def _format_size(size: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.1f}{unit}" if unit != "B" else f"{size}B"
        value /= 1024


def _read_xlsx_structure(raw: bytes) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    sheets = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            sheets.append({"name": sn, "columns": [], "total_rows": 0, "sample_rows": []})
            continue
        
        # 提取表头
        header = []
        for i, h in enumerate(all_rows[0]):
            header.append(str(h) if h is not None else f"col_{i}")
            
        # 样例行
        sample = []
        data_rows = [r for r in all_rows[1:] if any(v is not None for v in r)]
        if data_rows:
            sample.append(dict(zip(header, data_rows[0])))
            
        sheets.append({
            "name": sn,
            "columns": header,
            "total_rows": len(data_rows),
            "sample_rows": sample
        })
    wb.close()
    return {"sheets": sheets}


def _read_csv_structure(raw: bytes) -> dict:
    import csv
    text = raw.decode("utf-8-sig", errors="replace")
    
    # 自动探测分隔符
    delims = [",", ";", "\t", "|"]
    first_lines = [l for l in text.splitlines()[:5] if l.strip()]
    detected_delim = ","
    if len(first_lines) >= 2:
        for d in delims:
            counts = [l.count(d) for l in first_lines]
            if all(c > 0 for c in counts) and len(set(counts)) == 1:
                detected_delim = d
                break
                
    reader = csv.DictReader(io.StringIO(text), delimiter=detected_delim)
    rows = [dict(row) for row in reader]
    
    columns = list(rows[0].keys()) if rows else []
    sample_rows = rows[:1] if rows else []
    
    return {
        "delimiter": "Tab" if detected_delim == "\t" else detected_delim,
        "columns": columns,
        "total_rows": len(rows),
        "sample_rows": sample_rows
    }


def _read_pdf_structure(raw: bytes) -> dict:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        pages_count = len(pdf.pages)
        tables_count = 0
        text_snippets = []
        
        for i, page in enumerate(pdf.pages):
            tbls = page.extract_tables() or []
            tables_count += len(tbls)
            
            # 仅提取前 2 页文本大纲
            if i < 2:
                text = page.extract_text() or ""
                snippet = text.strip().replace("\n", " ")
                if len(snippet) > 200:
                    snippet = snippet[:200] + "..."
                text_snippets.append({
                    "page": i + 1,
                    "preview": snippet
                })
                
    return {
        "total_pages": pages_count,
        "total_tables_detected": tables_count,
        "preview_by_page": text_snippets
    }


def _read_docx_structure(raw: bytes) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    
    # 提取 Heading 样式的标题大纲
    headings = []
    for idx, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and p.style.name.startswith("Heading") and p.text.strip():
            headings.append({
                "paragraph_idx": idx + 1,
                "style": p.style.name,
                "text": p.text.strip()
            })
            
    # 提取前 3 个表格的维度和表头列名
    tables = []
    for idx, tbl in enumerate(doc.tables[:3]):
        rows_count = len(tbl.rows)
        cols_count = len(tbl.columns) if rows_count > 0 else 0
        headers = []
        if rows_count > 0:
            headers = [cell.text.strip() for cell in tbl.rows[0].cells]
            
        tables.append({
            "table_idx": idx + 1,
            "rows": rows_count,
            "columns_count": cols_count,
            "header_preview": headers
        })
        
    return {
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "headings_outline": headings,
        "tables_preview": tables
    }


def _read_json_structure(content: str) -> dict:
    data = json.loads(content)
    if isinstance(data, list):
        keys = []
        types = {}
        sample = None
        if data and isinstance(data[0], dict):
            sample = data[0]
            keys = list(sample.keys())
            types = {k: type(v).__name__ for k, v in sample.items()}
        return {
            "root_type": "array",
            "length": len(data),
            "keys": keys,
            "item_types": types,
            "sample_item": sample
        }
    elif isinstance(data, dict):
        keys = list(data.keys())
        types = {k: type(v).__name__ for k, v in data.items()}
        return {
            "root_type": "object",
            "keys": keys,
            "types": types,
            "keys_count": len(keys)
        }
    else:
        return {
            "root_type": type(data).__name__,
            "value_preview": str(data)[:100]
        }


def _read_md_structure(content: str) -> dict:
    lines = content.splitlines()
    headings = []
    tables = []
    
    # 提取 Heading 和 MD 表格
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 匹配标题
        if line.startswith("#"):
            m = re.match(r"^(#+)\s+(.*)$", line)
            if m:
                headings.append({
                    "line": i + 1,
                    "level": len(m.group(1)),
                    "text": m.group(2)
                })
        # 匹配 Markdown 表格
        elif line.startswith("|") and line.endswith("|") and i + 1 < len(lines):
            next_line = lines[i+1].strip()
            if next_line.startswith("|") and "-" in next_line:
                cols = [c.strip() for c in line.split("|")[1:-1]]
                table_rows = 0
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    table_rows += 1
                    j += 1
                tables.append({
                    "columns": cols,
                    "total_rows": table_rows,
                    "approx_line": i + 1
                })
                i = j - 1
        i += 1
        
    return {
        "headings_outline": headings,
        "detected_tables": tables
    }


def _read_txt_structure(content: str) -> dict:
    lines = content.splitlines()
    total_lines = len(lines)
    preview = [l for l in lines[:5]]
    
    # 自动探测伪 CSV/TSV
    delims = [",", "\t", "|", ";"]
    pseudo_csv = False
    detected_delim = None
    columns = []
    
    non_empty_lines = [l.strip() for l in lines[:5] if l.strip()]
    if len(non_empty_lines) >= 2:
        for d in delims:
            counts = [l.count(d) for l in non_empty_lines]
            if all(c > 0 for c in counts) and len(set(counts)) == 1:
                pseudo_csv = True
                detected_delim = d
                # 猜测首行为表头
                columns = [col.strip() for col in non_empty_lines[0].split(d)]
                break
                
    # 自动探测日志特征
    is_log = False
    log_keywords = {"info", "warn", "error", "debug", "fatal", "trace"}
    log_lines_matched = 0
    for l in lines[:10]:
        l_lower = l.lower()
        if any(kw in l_lower for kw in log_keywords) or re.search(r"\d{4}-\d{2}-\d{2}", l):
            log_lines_matched += 1
    if log_lines_matched >= 3:
        is_log = True
        
    structure_type = "plain_text"
    if pseudo_csv:
        structure_type = "pseudo_csv"
    elif is_log:
        structure_type = "log_file"
        
    summary = {
        "structure_type": structure_type,
        "total_lines": total_lines,
        "first_5_lines_preview": preview
    }
    
    if pseudo_csv:
        summary["pseudo_csv_details"] = {
            "delimiter": "Tab" if detected_delim == "\t" else detected_delim,
            "predicted_columns": columns
        }
    return summary


def _read_zip_structure(raw: bytes) -> dict:
    with zipfile.ZipFile(io.BytesIO(raw)) as z:
        infolist = z.infolist()
        files = []
        for info in infolist:
            files.append({
                "path": info.filename,
                "is_dir": info.is_dir(),
                "size": info.file_size,
                "size_human": _format_size(info.file_size)
            })
    return {
        "files_count": len(files),
        "files_list": files[:100]  # 最多显示前 100 个文件
    }


def _read_sqlite_structure(db_path: Path) -> dict:
    db_uri = f"file:{pathname2url(str(db_path.resolve()))}?mode=ro"
    conn = sqlite3.connect(db_uri, uri=True)
    try:
        cursor = conn.cursor()
        
        # 探测所有用户表
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
        tables = [r[0] for r in cursor.fetchall()]
        
        tables_summary = []
        for name in tables:
            # 探测列信息
            cursor.execute(f"PRAGMA table_info(\"{name}\")")
            cols = [col[1] for col in cursor.fetchall()]
            # 探测总行数
            cursor.execute(f"SELECT COUNT(*) FROM \"{name}\"")
            row_count = cursor.fetchone()[0]
            
            tables_summary.append({
                "table_name": name,
                "columns": cols,
                "total_rows": row_count
            })
            
        return {
            "tables_count": len(tables),
            "tables": tables_summary
        }
    finally:
        conn.close()


# ── 公开接口实现 ──

def read_document_structure_impl(
    ws: Workspace,
    path: str,
    emit_log: Optional[Callable[[str, str | dict], None]] = None,
) -> str:
    """探测指定文档的结构（元数据、列、大纲），绝不返回大文件全文，统一输出 JSON。"""
    try:
        p = ws.resolve_read(path)
    except ValueError as e:
        return json.dumps({"error": str(e), "path": path}, ensure_ascii=False)
    if p.name == "plan.json":
        return json.dumps({"error": f"不允许使用 read_document_structure 读取受限文件: {path}"}, ensure_ascii=False)
    if not p.exists():
        return json.dumps({"error": f"文件不存在: {path}"}, ensure_ascii=False)

    lower = path.lower()
    size = p.stat().st_size
    size_human = _format_size(size)

    # ZIP 和 SQLite 数据库不需要读全部 bytes
    if lower.endswith((".zip", ".tar", ".gz", ".rar")):
        try:
            raw = p.read_bytes()
            summary = _read_zip_structure(raw)
            kind = "zip"
            rec_tool = "run_python"
        except Exception as e:
            return json.dumps({"error": f"解析 ZIP 失败: {e}"}, ensure_ascii=False)
    elif lower.endswith((".sqlite", ".db")):
        try:
            summary = _read_sqlite_structure(p)
            kind = "sqlite"
            rec_tool = "query_sqlite"
        except Exception as e:
            return json.dumps({"error": f"解析 SQLite 失败: {e}"}, ensure_ascii=False)
    else:
        # 其他类型文本/二进制文件
        try:
            if lower.endswith((".xlsx", ".xls")):
                raw = p.read_bytes()
                summary = _read_xlsx_structure(raw)
                kind = "xlsx"
                rec_tool = "run_python"
            elif lower.endswith(".csv"):
                raw = p.read_bytes()
                summary = _read_csv_structure(raw)
                kind = "csv"
                rec_tool = "duckdb_register_parquet"
            elif lower.endswith(".pdf"):
                raw = p.read_bytes()
                summary = _read_pdf_structure(raw)
                kind = "pdf"
                rec_tool = "run_python"
            elif lower.endswith(".docx"):
                raw = p.read_bytes()
                summary = _read_docx_structure(raw)
                kind = "docx"
                rec_tool = "run_python"
            elif lower.endswith(".json"):
                content = p.read_text(encoding="utf-8", errors="replace")
                summary = _read_json_structure(content)
                kind = "json"
                rec_tool = "read_file" if size < 100 * 1024 else "run_python"
            elif lower.endswith(".md"):
                content = p.read_text(encoding="utf-8", errors="replace")
                summary = _read_md_structure(content)
                kind = "markdown"
                rec_tool = "read_file" if size < 200 * 1024 else "search_files"
            else:
                # 默认按普通 text 探测
                content = p.read_text(encoding="utf-8", errors="replace")
                summary = _read_txt_structure(content)
                kind = "text"
                rec_tool = "read_file" if size < 100 * 1024 else "search_files"
                if summary.get("structure_type") == "pseudo_csv":
                    rec_tool = "run_python"
        except Exception as e:
            return json.dumps({"error": f"读取文档结构失败: {e}"}, ensure_ascii=False)

    payload = {
        "path": path,
        "kind": kind,
        "size": size,
        "size_human": size_human,
        "summary": summary,
        "recommended_next_tool": rec_tool
    }
    result = json.dumps(payload, ensure_ascii=False)
    if emit_log:
        emit_log("custom_agent", {"level": "info", "message": f"✅ read_document_structure 调用成功: {path}"})
    return result


# 为向后兼容保留的别名接口
def read_document_impl(ws: Workspace, path: str) -> str:
    """向后兼容：调用最新的 read_document_structure_impl。"""
    return read_document_structure_impl(ws, path)


def _extract_xlsx(raw: bytes, sheet: str) -> list[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    sheet_names = wb.sheetnames
    if not sheet:
        target_sheets = [sheet_names[0]]
    else:
        target_sheets = [s for s in sheet_names if sheet.lower() in s.lower()]
        if not target_sheets:
            wb.close()
            raise ValueError(f"找不到 sheet '{sheet}'，可用: {sheet_names}")

    all_rows = []
    for sn in target_sheets:
        ws = wb[sn]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = []
        seen = set()
        for i, h in enumerate(rows[0]):
            base = str(h) if h is not None else f"col_{i}"
            name = base
            c = 1
            while name in seen:
                name = f"{base}_{c}"
                c += 1
            seen.add(name)
            header.append(name)
        data_rows = [dict(zip(header, r)) for r in rows[1:] if any(v is not None for v in r)]
        for r in data_rows:
            r["_sheet"] = sn
        all_rows.extend(data_rows)
    wb.close()
    return all_rows


def _extract_csv(raw: bytes) -> list[dict]:
    import csv
    text = raw.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    return [dict(row) for row in reader]


def _extract_pdf_tables(raw: bytes) -> list[dict]:
    import pdfplumber
    all_rows = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for pi, page in enumerate(pdf.pages):
            try:
                tables = page.extract_tables()
                for ti, tbl in enumerate(tables):
                    if not tbl or not tbl[0]:
                        continue
                    header = [str(c) if c else f"col_{ci}" for ci, c in enumerate(tbl[0])]
                    for row in tbl[1:]:
                        if any(c is not None for c in row):
                            row_dict = dict(zip(header, [str(c) if c else "" for c in row]))
                            row_dict["_page"] = pi + 1
                            row_dict["_table"] = ti + 1
                            all_rows.append(row_dict)
            except Exception:
                continue
    return all_rows


def _extract_docx_tables(raw: bytes) -> list[dict]:
    from docx import Document
    doc = Document(io.BytesIO(raw))
    all_rows = []
    for ti, tbl in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        if not rows:
            continue
        header = rows[0]
        for row in rows[1:]:
            row_dict = dict(zip(header, row))
            row_dict["_table"] = ti + 1
            all_rows.append(row_dict)
    return all_rows


def extract_document_tables_impl(ws: Workspace, path: str, sheet: str = "") -> str:
    """从 xlsx/csv/pdf/docx 提取结构化表格数据，返回 JSON 数组。"""
    try:
        p = ws.resolve_read(path)
    except ValueError as e:
        return json.dumps({"error": str(e), "path": path}, ensure_ascii=False)
    if not p.exists():
        return json.dumps({"error": f"文件不存在: {path}"})
    lower = path.lower()
    raw = p.read_bytes()
    try:
        if lower.endswith((".xlsx", ".xls")):
            rows = _extract_xlsx(raw, sheet)
        elif lower.endswith(".csv"):
            rows = _extract_csv(raw)
        elif lower.endswith(".pdf"):
            rows = _extract_pdf_tables(raw)
        elif lower.endswith(".docx"):
            rows = _extract_docx_tables(raw)
        else:
            return json.dumps({"error": f"不支持从此格式提取表格: {path}"})
        return json.dumps(rows, ensure_ascii=False, default=str)
    except Exception as e:
        return json.dumps({"error": f"提取表格失败: {e}"}, ensure_ascii=False)
